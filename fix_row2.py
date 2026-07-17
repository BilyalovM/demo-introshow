import docx
doc = docx.Document('templates/contract_template.docx')
table = doc.tables[3]
row = table.rows[1]
# Clear out all cells in this row to be safe
for cell in row.cells:
    for p in cell.paragraphs:
        p.text = p.text.replace('{% tr for item in items %}', '')
        p.text = p.text.replace('{%tr for item in items %}', '')
        p.text = p.text.replace('{% for item in items %}', '')
        p.text = p.text.replace('{% endfor %}', '')
        p.text = p.text.replace('{% tr endfor %}', '')
        p.text = p.text.replace('{%tr endfor %}', '')

row.cells[0].paragraphs[0].text = '{%tr for item in items %}' + row.cells[0].paragraphs[0].text
row.cells[-1].paragraphs[-1].text = row.cells[-1].paragraphs[-1].text + '{%tr endfor %}'

doc.save('templates/contract_template_fixed2.docx')

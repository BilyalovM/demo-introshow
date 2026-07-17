import docx
doc = docx.Document('templates/contract_template.docx')
table = doc.tables[3]
# we want to put {% tr for item in items %} at the very beginning of the row
row = table.rows[1] # The row we want to loop
row.cells[0].paragraphs[0].text = '{% tr for item in items %}' + row.cells[0].paragraphs[0].text.replace('{% for item in items %}', '').replace('{%tr for item in items %}', '')
row.cells[-1].paragraphs[-1].text = row.cells[-1].paragraphs[-1].text.replace('{% endfor %}', '').replace('{%tr endfor %}', '') + '{% tr endfor %}'
doc.save('templates/contract_template_fixed.docx')

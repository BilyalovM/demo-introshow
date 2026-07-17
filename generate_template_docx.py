from docx import Document
import os

def create_template():
    doc = Document()
    doc.add_heading('Договор аренды оборудования № {{ contract_number }}', 0)
    
    doc.add_paragraph('г. Алматы\t\t\t\tДата: {{ contract_date }}')
    
    doc.add_paragraph(
        'ИП "Rental Automation", именуемое в дальнейшем "Арендодатель", с одной стороны, и '
        '{{ company_name }}, в лице директора {{ director_name }}, именуемое в дальнейшем "Арендатор", '
        'заключили настоящий договор о нижеследующем:'
    )
    
    doc.add_heading('1. Предмет договора', level=1)
    doc.add_paragraph(
        'Арендодатель передает, а Арендатор принимает во временное пользование оборудование и услуги '
        'для проведения мероприятия "{{ event_name }}", которое состоится в период '
        '{{ event_date }} по адресу: {{ event_address }}.'
    )
    
    doc.add_heading('2. Стоимость аренды', level=1)
    doc.add_paragraph('Стоимость оборудования со скидкой ({{ discount_percentage }}%): {{ equipment_total }} тенге.')
    doc.add_paragraph('Стоимость логистики и персонала (фиксированная): {{ fixed_total }} тенге.')
    doc.add_paragraph('Итоговая сумма договора: {{ grand_total }} тенге.')
    doc.add_paragraph('Сумма прописью: {{ grand_total_text }}.')
    
    doc.add_heading('Приложение №1: Спецификация оборудования и услуг', level=1)
    
    # Table
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Наименование'
    hdr_cells[1].text = 'Категория'
    hdr_cells[2].text = 'Кол-во'
    hdr_cells[3].text = 'Дни'
    hdr_cells[4].text = 'Сумма со скидкой'
    
    # Jinja loop in docx table
    row_cells = table.rows[1].cells
    row_cells[0].text = '{% tr for item in items %}{{ item.name }}'
    row_cells[1].text = '{{ item.category_type }}'
    row_cells[2].text = '{{ item.quantity }}'
    row_cells[3].text = '{{ item.days }}'
    row_cells[4].text = '{{ item.line_total_discounted }}{% tr endfor %}'
    
    doc.add_heading('3. Реквизиты сторон', level=1)
    doc.add_paragraph('Арендатор:')
    doc.add_paragraph('Компания: {{ company_name }}')
    doc.add_paragraph('ИИН/БИН: {{ iin_bin }}')
    doc.add_paragraph('IBAN: {{ iban }}')
    
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/contract_template.docx')

if __name__ == '__main__':
    create_template()

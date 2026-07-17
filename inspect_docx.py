import docx

doc = docx.Document('templates/contract_template.docx')
for i, p in enumerate(doc.paragraphs):
    if 'ТЕСТ' in p.text or 'Тест' in p.text or 'Директор' in p.text or 'Шадрин' in p.text:
        print(f"Para {i}: {p.text}")

print("--- Tables ---")
for i, table in enumerate(doc.tables):
    print(f"Table {i}: {len(table.rows)} rows, {len(table.columns)} columns")
    for j, row in enumerate(table.rows):
        row_text = [cell.text.replace('\n', ' ') for cell in row.cells]
        print(f"  Row {j}: {row_text}")
        if j > 5:
            break

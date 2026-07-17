import docx
from docxtpl import DocxTemplate

doc = docx.Document()
table = doc.add_table(rows=3, cols=1)
table.cell(0, 0).text = "{%tr for item in items %}"
table.cell(1, 0).text = "Row Content: {{ item.name }}"
table.cell(2, 0).text = "{%tr endfor %}"

doc.save("test_table3.docx")

tpl = DocxTemplate("test_table3.docx")
context = {'items': [{'name': 'A'}, {'name': 'B'}]}
tpl.render(context)
tpl.save("test_table3_out.docx")

doc_out = docx.Document("test_table3_out.docx")
print(f"Result table rows: {len(doc_out.tables[0].rows)}")
for row in doc_out.tables[0].rows:
    print(row.cells[0].text)


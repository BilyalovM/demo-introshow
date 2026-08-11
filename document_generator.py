import os
import re
from collections import OrderedDict
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.text.paragraph import Paragraph
from num2words import num2words
from typing import Dict, Any, List, Optional, Tuple

# A4 (21.0 cm) − поля 2.0 cm слева/справа = 17.0 cm usable.
# Имя — flex: sum(фиксированных) + name ≤ usable. Не шире страницы (регресс c174c16: 16 cm
# при полях 2.54 cm + CENTER обрезало таблицу слева).
_DOCX_MARGIN_LR_CM = 2.0
_DOCX_USABLE_CM = 21.0 - 2 * _DOCX_MARGIN_LR_CM  # 17.0
_DOCX_SUM_W = 2.8
_DOCX_NO_W = 0.8
_DOCX_QTY_W = 1.5
_DOCX_SHIFTS_W = 1.5
_DOCX_PRICE_W = 2.2


def _docx_col_widths(fixed_after_name: List[float], *, no_w: float = _DOCX_NO_W) -> List[float]:
    """№ | name(flex) | …fixed… | сумма — name забирает остаток usable."""
    fixed_sum = no_w + sum(fixed_after_name)
    name_w = max(3.0, _DOCX_USABLE_CM - fixed_sum)
    return [no_w, name_w, *fixed_after_name]


_DOCX_W_TECH = _docx_col_widths([_DOCX_QTY_W + 0.5], no_w=0.9)  # № | name | Кол-во
_DOCX_W_6 = _docx_col_widths([_DOCX_QTY_W, _DOCX_PRICE_W, _DOCX_SHIFTS_W, _DOCX_SUM_W])
_DOCX_W_5 = _docx_col_widths([_DOCX_QTY_W, _DOCX_SHIFTS_W, _DOCX_SUM_W])  # клиент без цены
_DOCX_W_SUB = _docx_col_widths([2.0, 2.0, 1.8, 1.6, _DOCX_SUM_W], no_w=0.7)
_DOCX_W_TOTALS = [_DOCX_USABLE_CM - _DOCX_SUM_W, _DOCX_SUM_W]


def _apply_estimate_page(doc: Document) -> None:
    """Единые поля A4 — usable совпадает с _DOCX_USABLE_CM."""
    for section in doc.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(_DOCX_MARGIN_LR_CM)
        section.right_margin = Cm(_DOCX_MARGIN_LR_CM)
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)

# Метка склада в названии позиции — клиенту не показываем
_SUBRENTAL_NAME_RE = re.compile(r"\s*\(\s*субаренда\s*\)\s*", re.IGNORECASE)


def _client_display_name(name: Any) -> str:
    """Убирает «(субаренда)» из названия для клиентского экспорта."""
    text = str(name or "").strip()
    if not text:
        return ""
    cleaned = _SUBRENTAL_NAME_RE.sub(" ", text)
    return re.sub(r"\s{2,}", " ", cleaned).strip()


# Палитра КП IntroShow / Excel «Новый шаблон сметы.xlsx»
COLOR_CORAL = "F68560"       # оранжевые заголовки секций / итоги (как в КП)
COLOR_GRAY = "A5A5A5"        # строки ИТОГО, колонка кол-ва
COLOR_WHITE = "FFFFFF"
RGB_CORAL = (246, 133, 96)
RGB_GRAY = (165, 165, 165)

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_DEFAULT_LOGO = os.path.join(_BASE_DIR, "static", "img", "introshow_logo.png")
_FALLBACK_LOGO = os.path.join(_BASE_DIR, "static", "img", "logo.jpg")

# Категории логистики/персонала/расходников — не попадают в техничку
_FIXED_CATEGORY_NAMES = {
    "Логистика",
    "Персонал",
    "Расходники",
    "Логистика, Тех персонал",
    "Логистика/Тех персонал/Расходники",
}


def _is_fixed_category_name(category: Any) -> bool:
    cat = (category or "").strip()
    if not cat:
        return False
    if cat in _FIXED_CATEGORY_NAMES:
        return True
    low = cat.lower()
    return any(k in low for k in ("логистика", "персонал", "расходник"))


def filter_technichka_items(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Техничка = только оборудование склада/площадки, без логистики и персонала."""
    return [i for i in (items or []) if not _is_fixed_category_name(i.get("category"))]


def _resolve_logo_path(context: Optional[Dict[str, Any]] = None) -> Optional[str]:
    if context and (
        context.get("tpl_force_no_logo")
        or context.get("tpl_show_logo") is False
    ):
        return None
    candidates = []
    if context:
        candidates.append(context.get("logo_path"))
    candidates.extend((
        os.environ.get("INTROSHOW_LOGO_PATH"),
        _DEFAULT_LOGO,
        _FALLBACK_LOGO,
    ))
    for path in candidates:
        if path and os.path.isfile(path):
            return path
    return None


def _tpl_flag(context: Optional[Dict[str, Any]], key: str, default: bool = True) -> bool:
    if not context or key not in context:
        return default
    return bool(context.get(key))


def _add_plain_notes(doc: Document, text: str, *, size: int = 9, italic: bool = False) -> None:
    """Многострочные примечания из шаблона (после плейсхолдеров)."""
    if not (text or "").strip():
        return
    for line in str(text).splitlines():
        line = line.rstrip()
        if not line.strip():
            doc.add_paragraph("")
            continue
        p = doc.add_paragraph()
        r = p.add_run(line)
        if italic:
            r.italic = True
        _set_run_font(r, size=size)


def _resolve_doc_title(context: Dict[str, Any], default_title: str) -> str:
    custom = (context.get("tpl_custom_title") or "").strip()
    return custom or default_title


def resolve_photo_path(photo_url: str) -> str:
    if photo_url.startswith("/uploads/"):
        uploads_dir = os.environ.get("RENTAL_UPLOADS_DIR", "uploads")
        return os.path.join(uploads_dir, os.path.basename(photo_url))
    return photo_url.lstrip("/")


def get_rubles_text(amount: float) -> str:
    """
    Converts a number to Russian text.
    Handles the integer part and adds tiyn (decimals) for tenge.
    """
    try:
        integer_part = int(amount)
        decimal_part = int(round((amount - integer_part) * 100))

        text = num2words(integer_part, lang='ru')

        # format decimal part to always have 2 digits
        decimal_str = f"{decimal_part:02d}"

        return f"{text.capitalize()} тенге {decimal_str} тиын"
    except Exception as e:
        return f"{amount:,.2f} тенге"


def generate_contract(context: Dict[str, Any], template_path: str, output_path: str) -> str:
    """
    Generates a docx contract by injecting context into a template.
    Returns the path to the generated document.
    """
    doc = DocxTemplate(template_path)

    # Calculate the total in text
    total_cost = context.get('grand_total', 0.0)
    context['grand_total_text'] = get_rubles_text(total_cost)

    doc.render(context)
    doc.save(output_path)

    # Append appendix for photos and descriptions if they exist
    items = context.get('items', [])
    has_appendix = any(i.get('photo_url') or i.get('description') for i in items)

    if has_appendix:
        append_doc = Document(output_path)
        append_doc.add_page_break()
        append_doc.add_heading('Приложение: Спецификация оборудования', level=1)

        for item in items:
            if item.get('photo_url') or item.get('description'):
                append_doc.add_heading(item['name'], level=2)
                if item.get('description'):
                    append_doc.add_paragraph(item['description'])
                if item.get('photo_url'):
                    img_path = resolve_photo_path(item['photo_url'])
                    if os.path.exists(img_path):
                        append_doc.add_picture(img_path, width=Inches(3))

        append_doc.save(output_path)

    return output_path


def _fmt_money(v: float) -> str:
    return f"{v:,.0f}".replace(",", " ") + " ₸"


def _set_run_font(run, *, bold: bool = False, size: int = 11, color: Optional[Tuple[int, int, int]] = None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    if color:
        run.font.color.rgb = RGBColor(*color)


def _shade_cell(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'))


def _vcenter_cell(cell) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    size: int = 10,
    center: bool = False,
    align: Optional[str] = None,
) -> None:
    """align: 'left' | 'center' | 'right'. center=True — сокращение для align='center'."""
    # Полная очистка <w:p> (после merge иначе остаётся пустой абзац → лишний \\n)
    tc = cell._tc
    for el in tc.findall(qn("w:p")):
        tc.remove(el)
    p = OxmlElement("w:p")
    tc.append(p)
    para = Paragraph(p, cell)
    if align is None and center:
        align = "center"
    if align == "center":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif align == "left":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(str(text))
    _set_run_font(run, bold=bold, size=size)
    _vcenter_cell(cell)


def _set_table_col_widths(table, widths_cm: List[float]) -> None:
    """Фиксированные ширины колонок ≤ usable; LEFT + indent 0 (не уезжает за край)."""
    table.autofit = False
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total_cm = sum(widths_cm)
    if total_cm > _DOCX_USABLE_CM + 0.01:
        scale = _DOCX_USABLE_CM / total_cm
        widths_cm = [w * scale for w in widths_cm]
        total_cm = sum(widths_cm)
    total_twips = int(round(total_cm * 567))  # 1 cm ≈ 567 twips
    width_twips = [int(round(w * 567)) for w in widths_cm]
    # Поправка округления: сумма gridCol = tblW
    drift = total_twips - sum(width_twips)
    if width_twips:
        width_twips[1 if len(width_twips) > 1 else 0] += drift

    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for tag in ("w:tblW", "w:tblInd", "w:tblLayout"):
        for old in tblPr.findall(qn(tag)):
            tblPr.remove(old)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(total_twips))
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)
    tblLayout = OxmlElement("w:tblLayout")
    tblLayout.set(qn("w:type"), "fixed")
    tblPr.append(tblLayout)

    # tblGrid — Word надёжнее берёт ширины отсюда, чем только с tcW
    tblGrid = tbl.tblGrid
    if tblGrid is None:
        tblGrid = OxmlElement("w:tblGrid")
        tbl.insert(1, tblGrid)
    else:
        for old in list(tblGrid):
            tblGrid.remove(old)
    for tw in width_twips:
        gridCol = OxmlElement("w:gridCol")
        gridCol.set(qn("w:w"), str(tw))
        tblGrid.append(gridCol)

    for row in table.rows:
        seen_tc = set()
        idx = 0
        while idx < len(width_twips) and idx < len(row.cells):
            cell = row.cells[idx]
            tc = cell._tc
            if id(tc) in seen_tc:
                idx += 1
                continue
            seen_tc.add(id(tc))
            span = 1
            tcPr = tc.get_or_add_tcPr()
            gridSpan = tcPr.find(qn("w:gridSpan"))
            if gridSpan is not None:
                try:
                    span = max(1, int(gridSpan.get(qn("w:val"))))
                except (TypeError, ValueError):
                    span = 1
            tw = sum(width_twips[idx : idx + span])
            cell.width = Cm(tw / 567.0)
            for old in tcPr.findall(qn("w:tcW")):
                tcPr.remove(old)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:w"), str(tw))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            idx += span


def _item_section_name(item: Dict[str, Any]) -> str:
    cat = (item.get("category") or "").strip()
    return cat if cat else "Оборудование и услуги"


def _group_items(items: List[Dict[str, Any]]) -> "OrderedDict[str, List[Dict[str, Any]]]":
    grouped: "OrderedDict[str, List[Dict[str, Any]]]" = OrderedDict()
    for item in items:
        key = _item_section_name(item)
        grouped.setdefault(key, []).append(item)
    return grouped


def _shifts_label(context: Dict[str, Any]) -> str:
    shifts_label = context.get("shifts_label")
    if shifts_label is None and context.get("shifts") is not None:
        s = context["shifts"]
        try:
            sf = float(s)
            shifts_label = str(int(sf)) if sf == int(sf) else str(sf)
        except (TypeError, ValueError):
            shifts_label = str(s)
    return str(shifts_label) if shifts_label not in (None, "") else ""


def _event_dates_label(context: Dict[str, Any]) -> str:
    """Дата проведения — как в КП IntroShow (не складские выезд/возврат)."""
    explicit = (context.get("event_dates_label") or "").strip()
    if explicit:
        return explicit
    event_date = (context.get("event_date") or context.get("return_date") or "").strip()
    setup = (context.get("setup_date") or context.get("departure_date") or "").strip()
    if setup and event_date and setup != event_date:
        return f"{setup} — {event_date}"
    return event_date or setup or (context.get("rent_period") or "").strip()


def _meta_pairs(
    context: Dict[str, Any],
    *,
    with_assignee: bool = False,
    for_warehouse: bool = False,
    include_customer: bool = False,
) -> List[Tuple[str, str]]:
    """Поля шапки как в КП IntroShow: логотип слева, таблица справа."""
    meta_pairs: List[Tuple[str, str]] = []
    project = context.get("project_name") or context.get("event_name")
    if project:
        meta_pairs.append(("Наименование проекта", str(project)))
    if include_customer and context.get("company_name"):
        meta_pairs.append(("Клиент", str(context["company_name"])))
    if context.get("contact_name"):
        meta_pairs.append(("Контактное лицо проекта", str(context["contact_name"])))

    mgr = (
        context.get("project_manager_name")
        or context.get("manager_name")
        or context.get("sales_manager_name")
        or ""
    ).strip()
    phone = (
        context.get("manager_phone")
        or context.get("our_company_phone")
        or ""
    ).strip()
    mgr_val = " ".join(x for x in (mgr, phone) if x).strip()
    if mgr_val:
        meta_pairs.append(("менеджер проекта/телефон", mgr_val))

    city = (context.get("city") or "").strip()
    address = (context.get("event_address") or "").strip()
    loc = " / ".join(p for p in (city, address) if p)
    if loc:
        meta_pairs.append(("Город\\Локация", loc))

    event_dates = _event_dates_label(context)
    if event_dates:
        meta_pairs.append(("Дата проведения мероприятия", event_dates))

    shifts_label = _shifts_label(context)
    if shifts_label:
        meta_pairs.append(("Количество дней работы (смен)", shifts_label))

    if for_warehouse:
        depart = (context.get("departure_date") or "").strip()
        ret = (context.get("return_date") or "").strip()
        if depart or ret:
            meta_pairs.append(("Выезд оборудования со склада", depart or "—"))
            meta_pairs.append(("Возврат оборудования на склад", ret or "—"))

    if with_assignee and context.get("assignee_name"):
        meta_pairs.append(("Ответственный на объекте", str(context["assignee_name"])))
    return meta_pairs


def _add_letterhead(doc: Document, context: Dict[str, Any]) -> None:
    """Совместимость: делегирует в KP-шапку (логотип + мета-таблица)."""
    _add_kp_header(doc, context, with_assignee=False, for_warehouse=False, include_customer=False)


def _add_meta(doc: Document, context: Dict[str, Any], with_assignee: bool = False) -> None:
    """Совместимость: мета уже в KP-шапке; если вызывают отдельно — plain-список."""
    for label, value in _meta_pairs(context, with_assignee=with_assignee):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        _set_run_font(r1, bold=False, size=10)
        r2 = p.add_run(str(value))
        _set_run_font(r2, bold=True, size=10)


def _add_kp_header(
    doc: Document,
    context: Dict[str, Any],
    *,
    with_assignee: bool = False,
    for_warehouse: bool = False,
    include_customer: bool = False,
) -> None:
    """Шапка КП IntroShow: логотип слева, мета-таблица справа."""
    logo = _resolve_logo_path(context)
    pairs = _meta_pairs(
        context,
        with_assignee=with_assignee,
        for_warehouse=for_warehouse,
        include_customer=include_customer,
    )

    outer = doc.add_table(rows=1, cols=2)
    left, right = outer.rows[0].cells

    # --- left: logo + short company line ---
    for el in list(left._tc.findall(qn("w:p"))):
        left._tc.remove(el)
    p_logo = OxmlElement("w:p")
    left._tc.append(p_logo)
    para_logo = Paragraph(p_logo, left)
    if logo:
        run = para_logo.add_run()
        try:
            run.add_picture(logo, width=Cm(4.2))
        except Exception:
            logo = None
    if not logo:
        name = (context.get("our_company_name") or "Intro Show").strip()
        run = para_logo.add_run(name)
        _set_run_font(run, bold=True, size=14)

    company_bits = []
    if _tpl_flag(context, "tpl_show_company_block", True):
        phone = (context.get("our_company_phone") or "").strip()
        email = (context.get("our_company_email") or "").strip()
        if phone:
            company_bits.append(phone)
        if email:
            company_bits.append(email)
    if company_bits:
        p_c = left.add_paragraph()
        r = p_c.add_run(" · ".join(company_bits))
        _set_run_font(r, size=8)

    # --- right: 2-col meta table (вложенная через XML — у Cell нет add_table) ---
    for el in list(right._tc.findall(qn("w:p"))):
        right._tc.remove(el)
    if pairs:
        meta = doc.add_table(rows=len(pairs), cols=2)
        meta.style = "Table Grid"
        for i, (label, value) in enumerate(pairs):
            cells = meta.rows[i].cells
            _set_cell_text(cells[0], label, bold=False, size=8, align="left")
            _set_cell_text(cells[1], value, bold=True, size=8, align="left")
            _shade_cell(cells[0], "F2F2F2")
        _set_table_col_widths(meta, [5.2, 6.3])
        meta_tbl = meta._tbl
        meta_tbl.getparent().remove(meta_tbl)
        right._tc.append(meta_tbl)
    else:
        right.add_paragraph("")

    _set_table_col_widths(outer, [5.0, 12.0])
    # убрать рамку внешней таблицы
    tbl = outer._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)
    doc.add_paragraph("")


def _add_company_footer(doc: Document, context: Dict[str, Any]) -> None:
    """Реквизиты компании внизу (как в КП IntroShow)."""
    if not _tpl_flag(context, "tpl_include_company_contacts", True):
        return
    name = (context.get("our_company_name") or "").strip()
    phone = (context.get("our_company_phone") or "").strip()
    address = (context.get("our_company_address") or "").strip()
    email = (context.get("our_company_email") or "").strip()
    bin_code = (context.get("our_company_bin") or "").strip()
    if not any((name, phone, address, email, bin_code)):
        return
    doc.add_paragraph("")
    lines = []
    if name:
        lines.append(name)
    if address:
        lines.append(address)
    if phone:
        lines.append(f"тел. {phone}")
    if email:
        lines.append(email)
    if bin_code:
        lines.append(f"БИН {bin_code}")
    for line in lines:
        p = doc.add_paragraph()
        r = p.add_run(line)
        _set_run_font(r, size=9)


def _add_client_notes_and_signatures(doc: Document, context: Optional[Dict[str, Any]] = None) -> None:
    """Примечания и строки утверждения — как в клиентских КП."""
    doc.add_paragraph("")
    # Если в шаблоне заданы body_notes — дефолтные примечания не дублируем
    has_custom_body = bool((context or {}).get("tpl_body_notes"))
    if not has_custom_body:
        notes = [
            "Примечание: 1. Клиенту необходимо произвести 100% предоплату по данному "
            "коммерческому предложению до начала работ Компании по проекту.",
            "Примечание: 2. Расходные материалы не входят в стоимость оборудования, "
            "если не указано иное в смете.",
        ]
        for text in notes:
            p = doc.add_paragraph()
            r = p.add_run(text)
            _set_run_font(r, size=8)
    if not _tpl_flag(context, "tpl_include_signature", True):
        return
    doc.add_paragraph("")
    for label in (
        "Смету утвердил со стороны Исполнителя: _______________________________",
        "Смету утвердил со стороны Заказчика: _________________________________",
        "Дата утверждения сметы: ____________________",
    ):
        p = doc.add_paragraph()
        r = p.add_run(label)
        _set_run_font(r, size=9)


def _add_estimate_table(
    doc: Document,
    items: List[Dict[str, Any]],
    *,
    with_prices: bool = True,
    show_unit_price: bool = True,
    name_suffix_fn=None,
) -> None:
    """Таблица с coral-заголовками секций и серыми ИТОГО — как в Excel.

    with_prices=False — техничка (только № / название / кол-во).
    with_prices=True, show_unit_price=False — клиентская смета без колонки «Цена».
    """
    if not items:
        return

    if not with_prices:
        cols = 3
        headers = ["№", "Наименование", "Кол-во"]
        # header aligns: name left-ish via section title; qty center
        header_aligns = ["left", "center", "center"]
        col_widths = _DOCX_W_TECH
    elif show_unit_price:
        cols = 6
        headers = ["№", "Наименование", "Кол-во", "Цена за ед.", "Кол-во смен", "Сумма"]
        header_aligns = ["left", "center", "center", "right", "center", "right"]
        col_widths = _DOCX_W_6
    else:
        cols = 5
        headers = ["№", "Наименование", "Кол-во", "Смен", "Сумма"]
        header_aligns = ["left", "center", "center", "center", "right"]
        col_widths = _DOCX_W_5

    table = doc.add_table(rows=0, cols=cols)
    table.style = "Table Grid"

    global_idx = 0
    for section, section_items in _group_items(items).items():
        # Section / column header (coral): категория в №+Наименование (merge),
        # справа — Кол-во / Смен / Сумма (полное имя не обрезается).
        hdr_row = table.add_row()
        if cols >= 2:
            hdr_row.cells[0].merge(hdr_row.cells[1])
        # Важно: брать cells ПОСЛЕ merge (иначе устаревшая ссылка на удалённый tc)
        hdr = hdr_row.cells
        _set_cell_text(hdr[0], section, bold=True, size=10, align="left")
        for i in range(2, cols):
            _set_cell_text(hdr[i], headers[i], bold=True, size=9, align=header_aligns[i])
        for c in hdr:
            _shade_cell(c, COLOR_CORAL)

        section_sum = 0.0
        for item in section_items:
            global_idx += 1
            row = table.add_row().cells
            name = str(item.get("name", ""))
            if name_suffix_fn:
                extra = name_suffix_fn(item)
                if extra:
                    name = f"{name}{extra}"
            qty = item.get("quantity", 1)
            days = item.get("days", 1)
            line_total = float(item.get("line_total_discounted", item.get("line_total_base", 0)) or 0)
            section_sum += line_total

            _set_cell_text(row[0], str(global_idx), size=9, align="center")
            _set_cell_text(row[1], name, size=9, align="left")
            if with_prices and show_unit_price:
                _set_cell_text(row[2], str(qty), bold=True, size=9, align="center")
                _shade_cell(row[2], COLOR_GRAY)
                _set_cell_text(row[3], _fmt_money(item.get("price", 0)), bold=True, size=9, align="right")
                _set_cell_text(row[4], str(days), bold=True, size=9, align="center")
                _set_cell_text(row[5], _fmt_money(line_total), bold=True, size=9, align="right")
            elif with_prices:
                _set_cell_text(row[2], str(qty), bold=True, size=9, align="center")
                _shade_cell(row[2], COLOR_GRAY)
                _set_cell_text(row[3], str(days), bold=True, size=9, align="center")
                _set_cell_text(row[4], _fmt_money(line_total), bold=True, size=9, align="right")
            else:
                _set_cell_text(row[2], str(qty), bold=True, size=9, align="center")
                _shade_cell(row[2], COLOR_GRAY)

        # Section total (gray)
        if with_prices:
            tot = table.add_row().cells
            _set_cell_text(tot[0], "", size=9)
            _set_cell_text(tot[1], "ИТОГО KZT", bold=True, size=10, align="left")
            if show_unit_price:
                _set_cell_text(tot[2], "", size=9)
                _set_cell_text(tot[3], "", size=9)
                _set_cell_text(tot[4], "", size=9)
                _set_cell_text(tot[5], _fmt_money(section_sum), bold=True, size=10, align="right")
            else:
                _set_cell_text(tot[2], "", size=9)
                _set_cell_text(tot[3], "", size=9)
                _set_cell_text(tot[4], _fmt_money(section_sum), bold=True, size=10, align="right")
            for c in tot:
                _shade_cell(c, COLOR_GRAY)

    _set_table_col_widths(table, col_widths)


def _add_subrental_table(doc: Document, sub_items: List[Dict[str, Any]]) -> None:
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    headers = ["№", "Наименование", "Поставщик", "Цена клиенту", "Себест.", "Кол-во × смен", "Сумма"]
    header_aligns = ["center", "center", "center", "right", "right", "center", "right"]
    for i, htxt in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], htxt, bold=True, size=9, align=header_aligns[i])
        _shade_cell(table.rows[0].cells[i], COLOR_CORAL)
    for idx, item in enumerate(sub_items, 1):
        row = table.add_row().cells
        qty = item.get("quantity", 1)
        days = item.get("days", 1)
        _set_cell_text(row[0], str(idx), size=9, align="center")
        _set_cell_text(row[1], str(item.get("name", "")), size=9, align="left")
        _set_cell_text(row[2], str(item.get("supplier") or "—"), size=9, align="left")
        _set_cell_text(row[3], _fmt_money(item.get("price", 0)), size=9, align="right")
        _set_cell_text(row[4], _fmt_money(item.get("cost_price", 0)), size=9, align="right")
        _set_cell_text(row[5], f"{qty} × {days}", bold=True, size=9, align="center")
        _shade_cell(row[5], COLOR_GRAY)
        _set_cell_text(
            row[6],
            _fmt_money(item.get("line_total_discounted", item.get("line_total_base", 0))),
            bold=True,
            size=9,
            align="right",
        )
    _set_table_col_widths(table, _DOCX_W_SUB)


def _add_totals_block(doc: Document, context: Dict[str, Any], mode: str) -> None:
    disc_pct = float(context.get("discount_percentage") or 0)
    tax_pct = float(context.get("tax_percentage") or 0)
    eq_base = context.get("equipment_base")
    if eq_base is None:
        eq_base = context.get("equipment_total", 0)
    eq_total = context.get("equipment_total", 0)
    fixed_total = context.get("fixed_total", 0)
    discount_amount = context.get("discount_amount")
    if discount_amount is None and disc_pct:
        discount_amount = float(eq_base) - float(eq_total)
    after_discount = context.get("after_discount")
    if after_discount is None:
        after_discount = float(eq_total) + float(fixed_total)
    tax_amount = context.get("tax_amount") or 0
    grand = context.get("grand_total", 0)

    rows = [
        ("Итоговая сумма за оборудование", float(eq_base)),
    ]
    if disc_pct:
        rows.append((f"Скидка на оборудование {disc_pct:.0f}%", -float(discount_amount or 0)))
        rows.append(("Оборудование со скидкой", float(eq_total)))
    rows.append(("Работа персонала + расходники", float(fixed_total)))
    rows.append(("Сумма итого", float(after_discount)))
    if tax_pct or tax_amount:
        rows.append((f"Итоговая сумма за проект с учетом НДС {tax_pct:.0f}%", float(grand)))
    else:
        rows.append(("ИТОГО", float(grand)))

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        _set_cell_text(cells[0], label, bold=True, size=10, align="left")
        _set_cell_text(cells[1], _fmt_money(value), bold=True, size=10, align="right")
        _shade_cell(cells[0], COLOR_CORAL)
        _shade_cell(cells[1], COLOR_CORAL)
    _set_table_col_widths(table, _DOCX_W_TOTALS)

    p = doc.add_paragraph(get_rubles_text(float(grand or 0)))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if p.runs:
        _set_run_font(p.runs[0], bold=True, size=11)

    if mode == "internal":
        cost_total = float(context.get("cost_total") or 0)
        margin = context.get("margin")
        if margin is None:
            margin = float(grand or 0) - cost_total
        if cost_total or margin:
            doc.add_paragraph("")
            h = doc.add_paragraph()
            run = h.add_run("Маржа (внутренний блок)")
            _set_run_font(run, bold=True, size=11)
            mt = doc.add_table(rows=0, cols=2)
            mt.style = "Table Grid"
            for label, value in (
                ("Себестоимость субаренды", cost_total),
                ("Маржа", float(margin)),
            ):
                cells = mt.add_row().cells
                _set_cell_text(cells[0], label, bold=True, size=10, align="left")
                _set_cell_text(cells[1], _fmt_money(value), bold=True, size=10, align="right")
                _shade_cell(cells[0], COLOR_GRAY)
                _shade_cell(cells[1], COLOR_GRAY)
            _set_table_col_widths(mt, _DOCX_W_TOTALS)


def generate_technichka_docx(context: Dict[str, Any], output_path: str) -> str:
    """Техничка для склада: оборудование без цен; логистика/персонал исключены."""
    doc = Document()
    _apply_estimate_page(doc)
    _add_kp_header(
        doc,
        context,
        with_assignee=True,
        for_warehouse=True,
        include_customer=False,
    )

    default_title = f"ТЕХНИЧКА № {context.get('number', '')} от {context.get('date', '')}"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(_resolve_doc_title(context, default_title))
    _set_run_font(run, bold=True, size=14)

    _add_plain_notes(doc, context.get("tpl_body_notes") or "")

    if _tpl_flag(context, "tpl_include_items_table", True):
        items = filter_technichka_items(context.get("items", []))
        _add_estimate_table(doc, items, with_prices=False)

    note = doc.add_paragraph()
    nr = note.add_run(
        "Цены скрыты. Только оборудование для склада/площадки — "
        "без логистики и технического персонала."
    )
    nr.italic = True
    _set_run_font(nr, size=9)

    _add_plain_notes(doc, context.get("tpl_footer_notes") or "")
    if _tpl_flag(context, "tpl_include_company_contacts", False):
        _add_company_footer(doc, context)

    doc.save(output_path)
    return output_path


def generate_estimate_docx(
    context: Dict[str, Any],
    output_path: str,
    mode: str = "internal",
) -> str:
    """Генерирует смету (.docx).

    mode:
      - internal («Для нас»): свой склад в основной таблице, блок субаренды, себестоимость и маржа
      - client («Клиентская без цен»): позиции без колонки цены за ед., итог есть
      - client_priced («Клиентская с ценами»): цены за ед. видны, без маржи / себестоимости субаренды
    """
    mode = (mode or "internal").strip().lower()
    if mode not in ("internal", "client", "client_priced"):
        mode = "internal"
    is_client = mode in ("client", "client_priced")
    hide_sub = is_client or bool(context.get("hide_subrental_section"))
    show_unit_price = mode != "client"  # client_priced и internal — с ценами за ед.

    doc = Document()
    _apply_estimate_page(doc)
    _add_kp_header(
        doc,
        context,
        with_assignee=False,
        for_warehouse=(mode == "internal"),
        include_customer=(mode == "internal"),
    )

    if mode == "client":
        title_text = f"СМЕТА № {context.get('number', '')} от {context.get('date', '')} (без цен за ед.)"
    elif mode == "client_priced":
        title_text = f"СМЕТА № {context.get('number', '')} от {context.get('date', '')}"
    else:
        title_text = f"СМЕТА (ВНУТРЕННЯЯ) № {context.get('number', '')} от {context.get('date', '')}"
    title_text = _resolve_doc_title(context, title_text)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(title_text)
    _set_run_font(run, bold=True, size=14)

    _add_plain_notes(doc, context.get("tpl_body_notes") or "")

    all_items = list(context.get("items", []) or [])
    if hide_sub:
        # Клиенту: субаренда в основной таблице как обычные товары (без метки склада / «субаренда»)
        main_items = [
            {**i, "warehouse_type": "own", "name": _client_display_name(i.get("name", ""))}
            for i in all_items
        ]
        sub_items = []
    else:
        main_items = [i for i in all_items if (i.get("warehouse_type") or "own") != "subrental"]
        sub_items = [i for i in all_items if (i.get("warehouse_type") or "own") == "subrental"]
        if not main_items and not sub_items and all_items:
            main_items = all_items

    if _tpl_flag(context, "tpl_include_items_table", True):
        if main_items:
            _add_estimate_table(doc, main_items, with_prices=True, show_unit_price=show_unit_price)
        elif is_client:
            doc.add_paragraph("Нет позиций для клиентской сметы.")

        if mode == "internal" and sub_items:
            doc.add_paragraph("")
            h = doc.add_paragraph()
            run = h.add_run("Субаренда (только для нас)")
            _set_run_font(run, bold=True, size=12)
            note = doc.add_paragraph()
            nr = note.add_run("В клиентской смете эти позиции идут как обычные, без себестоимости.")
            nr.italic = True
            _add_subrental_table(doc, sub_items)

    if _tpl_flag(context, "tpl_include_totals", True):
        doc.add_paragraph("")
        _add_totals_block(doc, context, mode)

    if is_client:
        _add_client_notes_and_signatures(doc, context)
        _add_company_footer(doc, context)
    elif mode == "internal":
        _add_company_footer(doc, context)

    _add_plain_notes(doc, context.get("tpl_footer_notes") or "")

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# PDF (fpdf2) — без системных зависимостей, работает на Vercel / Linux
# ---------------------------------------------------------------------------

_FONTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "static", "fonts")


def _pdf_font_paths():
    regular = os.path.join(_FONTS_DIR, "DejaVuSans.ttf")
    bold = os.path.join(_FONTS_DIR, "DejaVuSans-Bold.ttf")
    return regular, bold


class _EstimatePDF:
    """Минимальная обёртка над FPDF с кириллицей (DejaVu) и палитрой Excel."""

    def __init__(self):
        from fpdf import FPDF

        self.pdf = FPDF(orientation="P", unit="mm", format="A4")
        self.pdf.set_auto_page_break(auto=True, margin=15)
        regular, bold = _pdf_font_paths()
        if not os.path.exists(regular):
            raise FileNotFoundError(f"PDF font missing: {regular}")
        self.pdf.add_font("DejaVu", "", regular)
        if os.path.exists(bold):
            self.pdf.add_font("DejaVu", "B", bold)
        else:
            self.pdf.add_font("DejaVu", "B", regular)
        self.pdf.add_page()
        self.pdf.set_font("DejaVu", size=10)

    def _reset_x(self):
        self.pdf.set_x(self.pdf.l_margin)

    def title(self, text: str):
        self._reset_x()
        self.pdf.set_font("DejaVu", "B", 14)
        self.pdf.multi_cell(0, 8, text, align="C")
        self._reset_x()
        self.pdf.ln(2)
        self.pdf.set_font("DejaVu", size=10)

    def line(self, text: str, bold: bool = False):
        self._reset_x()
        self.pdf.set_font("DejaVu", "B" if bold else "", 10)
        self.pdf.multi_cell(0, 5, text)
        self._reset_x()
        self.pdf.set_font("DejaVu", size=10)

    def right(self, text: str, bold: bool = False, size: int = 10):
        self._reset_x()
        self.pdf.set_font("DejaVu", "B" if bold else "", size)
        self.pdf.cell(0, 6, text, align="R", new_x="LMARGIN", new_y="NEXT")
        self.pdf.set_font("DejaVu", size=10)

    def table(
        self,
        headers: List[str],
        rows: List[List[str]],
        col_widths: List[float],
        *,
        header_rgb: Tuple[int, int, int] = RGB_CORAL,
        fill_header: bool = True,
        aligns: Optional[List[str]] = None,
    ):
        self._reset_x()
        usable = self.pdf.w - self.pdf.l_margin - self.pdf.r_margin
        total_w = sum(col_widths)
        if total_w > usable and total_w > 0:
            scale = usable / total_w
            col_widths = [w * scale for w in col_widths]
        n = len(col_widths)
        aligns = list(aligns or ["L"] * n)
        while len(aligns) < n:
            aligns.append("L")

        if fill_header:
            self.pdf.set_fill_color(*header_rgb)
            self.pdf.set_font("DejaVu", "B", 8)
            for i, h in enumerate(headers):
                self.pdf.cell(col_widths[i], 6, h[:40], border=1, fill=True, align=aligns[i])
            self.pdf.ln()
            self._reset_x()

        self.pdf.set_font("DejaVu", size=8)
        for row in rows:
            cells = [str(c)[:80] for c in row]
            row_h = 6
            x0 = self.pdf.l_margin
            y0 = self.pdf.get_y()
            if y0 + row_h > self.pdf.h - 15:
                self.pdf.add_page()
                self._reset_x()
                y0 = self.pdf.get_y()
            is_total = cells and ("ИТОГО" in cells[0].upper() or "ИТОГО" in (cells[1].upper() if len(cells) > 1 else ""))
            if is_total:
                self.pdf.set_fill_color(*RGB_GRAY)
                self.pdf.set_font("DejaVu", "B", 8)
            for i, cell in enumerate(cells):
                x = x0 + sum(col_widths[:i])
                self.pdf.set_xy(x, y0)
                self.pdf.cell(col_widths[i], row_h, cell, border=1, fill=is_total, align=aligns[i])
            if is_total:
                self.pdf.set_font("DejaVu", size=8)
            self.pdf.set_xy(x0, y0 + row_h)
        self._reset_x()
        self.pdf.set_font("DejaVu", size=10)

    def totals_table(self, rows: List[Tuple[str, str]], sum_col_w: Optional[float] = None):
        """Coral totals: full page width, amounts right-aligned in Сумма-width column."""
        self._reset_x()
        usable = self.pdf.w - self.pdf.l_margin - self.pdf.r_margin
        w_val = float(sum_col_w) if sum_col_w else min(38.0, usable * 0.22)
        w_label = usable - w_val
        self.pdf.set_fill_color(*RGB_CORAL)
        self.pdf.set_font("DejaVu", "B", 9)
        for label, value in rows:
            y0 = self.pdf.get_y()
            if y0 + 7 > self.pdf.h - 15:
                self.pdf.add_page()
                self._reset_x()
            self.pdf.cell(w_label, 7, label[:70], border=1, fill=True, align="L")
            self.pdf.cell(w_val, 7, value[:40], border=1, fill=True, align="R")
            self.pdf.ln()
            self._reset_x()
        self.pdf.set_font("DejaVu", size=10)

    def save(self, path: str):
        self.pdf.output(path)


def _pdf_letterhead_lines(context: Dict[str, Any]) -> List[str]:
    if not _tpl_flag(context, "tpl_include_company_contacts", True):
        return []
    name = (context.get("our_company_name") or "").strip()
    phone = (context.get("our_company_phone") or "").strip()
    email = (context.get("our_company_email") or "").strip()
    address = (context.get("our_company_address") or "").strip()
    bin_code = (context.get("our_company_bin") or "").strip()
    lines = []
    if name:
        lines.append(name)
    if address:
        lines.append(address)
    contact_bits = [b for b in (phone, email) if b]
    if contact_bits:
        lines.append(" · ".join(contact_bits))
    if bin_code:
        lines.append(f"БИН {bin_code}")
    return lines


def _pdf_add_notes(pdf: "_EstimatePDF", text: str) -> None:
    if not (text or "").strip():
        return
    pdf.pdf.ln(1)
    for line in str(text).splitlines():
        if not line.strip():
            pdf.pdf.ln(2)
            continue
        pdf.line(line)


def _pdf_meta_lines(
    context: Dict[str, Any],
    with_assignee: bool = False,
    *,
    for_warehouse: bool = False,
    include_customer: bool = False,
) -> List[str]:
    return [
        f"{label}: {value}"
        for label, value in _meta_pairs(
            context,
            with_assignee=with_assignee,
            for_warehouse=for_warehouse,
            include_customer=include_customer,
        )
    ]


def _pdf_add_kp_header(
    pdf: "_EstimatePDF",
    context: Dict[str, Any],
    *,
    with_assignee: bool = False,
    for_warehouse: bool = False,
    include_customer: bool = False,
) -> None:
    """PDF-шапка: логотип слева, мета-таблица справа (как КП IntroShow)."""
    logo = _resolve_logo_path(context)
    pairs = _meta_pairs(
        context,
        with_assignee=with_assignee,
        for_warehouse=for_warehouse,
        include_customer=include_customer,
    )
    usable = pdf.pdf.w - pdf.pdf.l_margin - pdf.pdf.r_margin
    logo_w = 42.0
    gap = 4.0
    meta_w = usable - logo_w - gap
    label_w = meta_w * 0.48
    value_w = meta_w - label_w
    x0 = pdf.pdf.l_margin
    y0 = pdf.pdf.get_y()

    logo_h = 0.0
    if logo:
        try:
            pdf.pdf.image(logo, x=x0, y=y0, w=logo_w)
            logo_h = 14.0
        except Exception:
            logo = None
    if not logo:
        pdf.pdf.set_xy(x0, y0)
        pdf.pdf.set_font("DejaVu", "B", 12)
        name = (context.get("our_company_name") or "Intro Show").strip()
        pdf.pdf.multi_cell(logo_w, 6, name)
        logo_h = max(logo_h, pdf.pdf.get_y() - y0)

    if _tpl_flag(context, "tpl_show_company_block", True):
        phone = (context.get("our_company_phone") or "").strip()
        email = (context.get("our_company_email") or "").strip()
        contact = " · ".join(x for x in (phone, email) if x)
        if contact:
            pdf.pdf.set_xy(x0, y0 + logo_h + 1)
            pdf.pdf.set_font("DejaVu", "", 7)
            pdf.pdf.multi_cell(logo_w, 3.5, contact[:80])
            logo_h = pdf.pdf.get_y() - y0

    meta_h = 0.0
    if pairs:
        pdf.pdf.set_font("DejaVu", size=8)
        y = y0
        for label, value in pairs:
            pdf.pdf.set_xy(x0 + logo_w + gap, y)
            pdf.pdf.set_fill_color(242, 242, 242)
            pdf.pdf.set_font("DejaVu", "", 7)
            pdf.pdf.cell(label_w, 5, str(label)[:40], border=1, fill=True, align="L")
            pdf.pdf.set_font("DejaVu", "B", 7)
            pdf.pdf.cell(value_w, 5, str(value)[:55], border=1, fill=False, align="L")
            y += 5
        meta_h = y - y0

    pdf.pdf.set_y(y0 + max(logo_h, meta_h) + 3)
    pdf.pdf.set_font("DejaVu", size=10)
    pdf._reset_x()


# PDF: единая ширина колонки «Сумма» (мм) — совпадает с правым краем блока итогов
_PDF_SUM_W = 34.0


def _pdf_col_aligns(*, with_prices: bool, show_unit_price: bool) -> List[str]:
    """L/C/R для fpdf cell align."""
    if not with_prices:
        return ["C", "L", "C"]
    if show_unit_price:
        return ["C", "L", "C", "R", "C", "R"]
    return ["C", "L", "C", "C", "R"]


def _pdf_build_sectioned_rows(
    items: List[Dict[str, Any]],
    *,
    with_prices: bool = True,
    show_unit_price: bool = True,
) -> Tuple[List[str], List[List[str]], List[float]]:
    if not with_prices:
        headers = ["№", "Наименование", "Кол-во"]
        widths = [12, 134, 28]
    elif show_unit_price:
        headers = ["№", "Наименование", "Кол-во", "Цена", "Смен", "Сумма"]
        widths = [10, 68, 18, 28, 18, _PDF_SUM_W]
    else:
        headers = ["№", "Наименование", "Кол-во", "Смен", "Сумма"]
        widths = [10, 88, 20, 20, _PDF_SUM_W]

    rows: List[List[str]] = []
    idx = 0
    for section, section_items in _group_items(items).items():
        if not with_prices:
            rows.append([section, "", "Кол-во"])
        elif show_unit_price:
            rows.append([section, "", "Кол-во", "Цена", "Смен", "Сумма"])
        else:
            rows.append([section, "", "Кол-во", "Смен", "Сумма"])
        # Mark section header by prefix for fill detection — use coral via special marker
        rows[-1][0] = f"§ {section}"
        section_sum = 0.0
        for item in section_items:
            idx += 1
            qty = item.get("quantity", 1)
            days = item.get("days", 1)
            line_total = float(item.get("line_total_discounted", item.get("line_total_base", 0)) or 0)
            section_sum += line_total
            if with_prices and show_unit_price:
                rows.append([
                    str(idx),
                    str(item.get("name", ""))[:60],
                    str(qty),
                    _fmt_money(item.get("price", 0)),
                    str(days),
                    _fmt_money(line_total),
                ])
            elif with_prices:
                rows.append([
                    str(idx),
                    str(item.get("name", ""))[:70],
                    str(qty),
                    str(days),
                    _fmt_money(line_total),
                ])
            else:
                rows.append([str(idx), str(item.get("name", ""))[:70], str(qty)])
        if with_prices and show_unit_price:
            rows.append(["", "ИТОГО KZT", "", "", "", _fmt_money(section_sum)])
        elif with_prices:
            rows.append(["", "ИТОГО KZT", "", "", _fmt_money(section_sum)])
    return headers, rows, widths


def _pdf_draw_sectioned_table(
    pdf: _EstimatePDF,
    items: List[Dict[str, Any]],
    *,
    with_prices: bool = True,
    show_unit_price: bool = True,
) -> float:
    """Draw grouped table with coral section headers and gray totals.

    Returns scaled width of the last column (Сумма / Кол-во) for totals alignment.
    """
    if not items:
        return _PDF_SUM_W
    headers, rows, col_widths = _pdf_build_sectioned_rows(
        items, with_prices=with_prices, show_unit_price=show_unit_price
    )
    usable = pdf.pdf.w - pdf.pdf.l_margin - pdf.pdf.r_margin
    total_w = sum(col_widths)
    if total_w > usable and total_w > 0:
        scale = usable / total_w
        col_widths = [w * scale for w in col_widths]
    aligns = _pdf_col_aligns(with_prices=with_prices, show_unit_price=show_unit_price)

    # Skip generic header — each section has its own coral header row
    pdf.pdf.set_font("DejaVu", size=8)
    for row in rows:
        cells = [str(c)[:80] for c in row]
        row_h = 6
        x0 = pdf.pdf.l_margin
        y0 = pdf.pdf.get_y()
        if y0 + row_h > pdf.pdf.h - 15:
            pdf.pdf.add_page()
            pdf._reset_x()
            y0 = pdf.pdf.get_y()

        is_section = cells[0].startswith("§ ")
        is_total = (not is_section) and any("ИТОГО" in c.upper() for c in cells)

        if is_section:
            pdf.pdf.set_fill_color(*RGB_CORAL)
            pdf.pdf.set_font("DejaVu", "B", 8)
            # Категория на ширине №+Наименование; справа — заголовки колонок
            label = cells[0][2:]  # strip §
            span_w = col_widths[0] + (col_widths[1] if len(col_widths) > 1 else 0)
            pdf.pdf.set_xy(x0, y0)
            pdf.pdf.cell(span_w, row_h, label[:70], border=1, fill=True, align="L")
            x = x0 + span_w
            for i in range(2, len(cells)):
                pdf.pdf.set_xy(x, y0)
                a = aligns[i] if i < len(aligns) else "L"
                pdf.pdf.cell(col_widths[i], row_h, cells[i][:40], border=1, fill=True, align=a)
                x += col_widths[i]
            pdf.pdf.set_font("DejaVu", size=8)
        elif is_total:
            pdf.pdf.set_fill_color(*RGB_GRAY)
            pdf.pdf.set_font("DejaVu", "B", 8)
            for i, cell in enumerate(cells):
                x = x0 + sum(col_widths[:i])
                pdf.pdf.set_xy(x, y0)
                a = aligns[i] if i < len(aligns) else "L"
                pdf.pdf.cell(col_widths[i], row_h, cell, border=1, fill=True, align=a)
            pdf.pdf.set_font("DejaVu", size=8)
        else:
            for i, cell in enumerate(cells):
                x = x0 + sum(col_widths[:i])
                pdf.pdf.set_xy(x, y0)
                # Gray qty column like Excel
                fill_qty = with_prices and i == 2
                if fill_qty:
                    pdf.pdf.set_fill_color(*RGB_GRAY)
                a = aligns[i] if i < len(aligns) else "L"
                pdf.pdf.cell(col_widths[i], row_h, cell, border=1, fill=fill_qty, align=a)
        pdf.pdf.set_xy(x0, y0 + row_h)
    pdf._reset_x()
    pdf.pdf.set_font("DejaVu", size=10)
    return col_widths[-1]


def generate_estimate_pdf(
    context: Dict[str, Any],
    output_path: str,
    mode: str = "internal",
) -> str:
    """PDF-смета: те же mode=internal|client|client_priced и данные, что у DOCX."""
    mode = (mode or "internal").strip().lower()
    if mode not in ("internal", "client", "client_priced"):
        mode = "internal"
    is_client = mode in ("client", "client_priced")
    show_unit_price = mode != "client"

    pdf = _EstimatePDF()
    _pdf_add_kp_header(
        pdf,
        context,
        with_assignee=False,
        for_warehouse=(mode == "internal"),
        include_customer=(mode == "internal"),
    )
    if mode == "client":
        title_text = f"СМЕТА № {context.get('number', '')} от {context.get('date', '')} (без цен за ед.)"
    elif mode == "client_priced":
        title_text = f"СМЕТА № {context.get('number', '')} от {context.get('date', '')}"
    else:
        title_text = f"СМЕТА (ВНУТРЕННЯЯ) № {context.get('number', '')} от {context.get('date', '')}"
    pdf.title(_resolve_doc_title(context, title_text))
    pdf.pdf.ln(1)
    _pdf_add_notes(pdf, context.get("tpl_body_notes") or "")

    all_items = list(context.get("items", []) or [])
    hide_sub = is_client or bool(context.get("hide_subrental_section"))
    if hide_sub:
        main_items = [
            {**i, "warehouse_type": "own", "name": _client_display_name(i.get("name", ""))}
            for i in all_items
        ]
        sub_items = []
    else:
        main_items = [i for i in all_items if (i.get("warehouse_type") or "own") != "subrental"]
        sub_items = [i for i in all_items if (i.get("warehouse_type") or "own") == "subrental"]
        if not main_items and not sub_items and all_items:
            main_items = all_items

    sum_col_w = _PDF_SUM_W
    if _tpl_flag(context, "tpl_include_items_table", True):
        if main_items:
            sum_col_w = _pdf_draw_sectioned_table(
                pdf, main_items, with_prices=True, show_unit_price=show_unit_price
            )
        elif is_client:
            pdf.line("Нет позиций для клиентской сметы.")

        if mode == "internal" and sub_items:
            pdf.pdf.ln(3)
            pdf.line("Субаренда (только для нас)", bold=True)
            pdf.line("В клиентской смете эти позиции идут как обычные, без себестоимости.")
            rows = []
            for idx, item in enumerate(sub_items, 1):
                qty = item.get("quantity", 1)
                days = item.get("days", 1)
                rows.append([
                    str(idx),
                    str(item.get("name", ""))[:40],
                    str(item.get("supplier") or "—")[:20],
                    _fmt_money(item.get("price", 0)),
                    _fmt_money(item.get("cost_price", 0)),
                    f"{qty}×{days}",
                    _fmt_money(item.get("line_total_discounted", item.get("line_total_base", 0))),
                ])
            sub_widths = [8, 46, 28, 24, 24, 16, _PDF_SUM_W]
            usable = pdf.pdf.w - pdf.pdf.l_margin - pdf.pdf.r_margin
            if sum(sub_widths) > usable:
                scale = usable / sum(sub_widths)
                sub_widths = [w * scale for w in sub_widths]
                sum_col_w = sub_widths[-1]
            pdf.table(
                ["№", "Наименование", "Поставщик", "Цена", "Себест.", "К×С", "Сумма"],
                rows,
                sub_widths,
                aligns=["C", "L", "L", "R", "R", "C", "R"],
            )

    disc_pct = float(context.get("discount_percentage") or 0)
    tax_pct = float(context.get("tax_percentage") or 0)
    eq_base = context.get("equipment_base")
    if eq_base is None:
        eq_base = context.get("equipment_total", 0)
    eq_total = context.get("equipment_total", 0)
    fixed_total = context.get("fixed_total", 0)
    discount_amount = context.get("discount_amount")
    if discount_amount is None and disc_pct:
        discount_amount = float(eq_base) - float(eq_total)
    after_discount = context.get("after_discount")
    if after_discount is None:
        after_discount = float(eq_total) + float(fixed_total)
    tax_amount = context.get("tax_amount") or 0
    grand = context.get("grand_total", 0)

    if _tpl_flag(context, "tpl_include_totals", True):
        totals_rows = [("Итоговая сумма за оборудование", _fmt_money(float(eq_base)))]
        if disc_pct:
            totals_rows.append((f"Скидка на оборудование {disc_pct:.0f}%", f"−{_fmt_money(float(discount_amount or 0))}"))
            totals_rows.append(("Оборудование со скидкой", _fmt_money(float(eq_total))))
        totals_rows.append(("Работа персонала + расходники", _fmt_money(float(fixed_total))))
        totals_rows.append(("Сумма итого", _fmt_money(float(after_discount))))
        if tax_pct or tax_amount:
            totals_rows.append(
                (f"Итоговая сумма за проект с учетом НДС {tax_pct:.0f}%", _fmt_money(float(grand)))
            )
        else:
            totals_rows.append(("ИТОГО", _fmt_money(float(grand))))

        pdf.pdf.ln(3)
        pdf.totals_table(totals_rows, sum_col_w=sum_col_w)
        pdf.right(get_rubles_text(float(grand or 0)), bold=True)

        if mode == "internal":
            cost_total = float(context.get("cost_total") or 0)
            margin = context.get("margin")
            if margin is None:
                margin = float(grand or 0) - cost_total
            if cost_total or margin:
                pdf.pdf.ln(2)
                pdf.line("Маржа (внутренний блок)", bold=True)
                pdf.pdf.set_fill_color(*RGB_GRAY)
                usable = pdf.pdf.w - pdf.pdf.l_margin - pdf.pdf.r_margin
                w_val = float(sum_col_w)
                w_label = usable - w_val
                pdf.pdf.set_font("DejaVu", "B", 9)
                for label, value in (
                    ("Себестоимость субаренды", _fmt_money(cost_total)),
                    ("Маржа", _fmt_money(float(margin))),
                ):
                    pdf.pdf.cell(w_label, 7, label, border=1, fill=True, align="L")
                    pdf.pdf.cell(w_val, 7, value, border=1, fill=True, align="R")
                    pdf.pdf.ln()
                    pdf._reset_x()
                pdf.pdf.set_font("DejaVu", size=10)

    if is_client:
        pdf.pdf.ln(3)
        has_custom_body = bool(context.get("tpl_body_notes"))
        if not has_custom_body:
            pdf.pdf.set_font("DejaVu", size=7)
            for note in (
                "Примечание: 1. Клиенту необходимо произвести 100% предоплату по данному "
                "коммерческому предложению до начала работ Компании по проекту.",
                "Примечание: 2. Расходные материалы не входят в стоимость оборудования, "
                "если не указано иное в смете.",
            ):
                pdf.line(note)
        if _tpl_flag(context, "tpl_include_signature", True):
            pdf.pdf.ln(2)
            pdf.pdf.set_font("DejaVu", size=9)
            pdf.line("Смету утвердил со стороны Исполнителя: _______________________________")
            pdf.line("Смету утвердил со стороны Заказчика: _________________________________")
            pdf.line("Дата утверждения сметы: ____________________")
        for lh in _pdf_letterhead_lines(context):
            pdf.line(lh)
    else:
        for lh in _pdf_letterhead_lines(context):
            pdf.line(lh)

    _pdf_add_notes(pdf, context.get("tpl_footer_notes") or "")

    pdf.save(output_path)
    return output_path


def generate_technichka_pdf(context: Dict[str, Any], output_path: str) -> str:
    """PDF-техничка: только оборудование, без логистики/персонала."""
    pdf = _EstimatePDF()
    _pdf_add_kp_header(
        pdf,
        context,
        with_assignee=True,
        for_warehouse=True,
        include_customer=False,
    )
    default_title = f"ТЕХНИЧКА № {context.get('number', '')} от {context.get('date', '')}"
    pdf.title(_resolve_doc_title(context, default_title))
    pdf.pdf.ln(1)
    _pdf_add_notes(pdf, context.get("tpl_body_notes") or "")
    if _tpl_flag(context, "tpl_include_items_table", True):
        items = filter_technichka_items(list(context.get("items", []) or []))
        _pdf_draw_sectioned_table(pdf, items, with_prices=False)
    pdf.pdf.ln(2)
    pdf.line(
        "Цены скрыты. Только оборудование для склада/площадки — "
        "без логистики и технического персонала."
    )
    _pdf_add_notes(pdf, context.get("tpl_footer_notes") or "")
    for lh in _pdf_letterhead_lines(context):
        pdf.line(lh)
    pdf.save(output_path)
    return output_path


def generate_contract_pdf(context: Dict[str, Any], output_path: str) -> str:
    """Упрощённый PDF договора: реквизиты + спецификация (как приложение к Word-шаблону)."""
    pdf = _EstimatePDF()
    _pdf_add_kp_header(
        pdf,
        context,
        with_assignee=False,
        for_warehouse=False,
        include_customer=True,
    )
    default_title = (
        f"ДОГОВОР № {context.get('contract_number', '')} от {context.get('contract_date', '')}"
    )
    pdf.title(_resolve_doc_title(context, default_title))
    _pdf_add_notes(pdf, context.get("tpl_body_notes") or "")
    pdf.line(f"Заказчик: {context.get('company_name') or '—'}", bold=True)
    if context.get("director_name"):
        pdf.line(f"Директор / представитель: {context['director_name']}")
    if context.get("iin_bin"):
        pdf.line(f"ИИН/БИН: {context['iin_bin']}")
    if context.get("iban"):
        pdf.line(f"ИБан / реквизиты: {context['iban']}")
    if context.get("event_name"):
        pdf.line(f"Мероприятие: {context['event_name']}")
    if context.get("event_date"):
        pdf.line(f"Дата: {context['event_date']}")
    if context.get("event_address"):
        pdf.line(f"Адрес: {context['event_address']}")
    pdf.pdf.ln(2)

    items = list(context.get("items", []) or [])
    sum_col_w = _PDF_SUM_W
    if _tpl_flag(context, "tpl_include_items_table", True):
        pdf.line("Спецификация оборудования и услуг", bold=True)
        if items:
            sum_col_w = _pdf_draw_sectioned_table(pdf, items, with_prices=True)
        else:
            pdf.line("Нет позиций.")

    if _tpl_flag(context, "tpl_include_totals", True):
        disc_pct = float(context.get("discount_percentage") or 0)
        tax_pct = float(context.get("tax_percentage") or 0)
        grand = float(context.get("grand_total") or 0)
        totals = [
            ("Оборудование", _fmt_money(float(context.get("equipment_total") or 0))),
            ("Логистика и персонал", _fmt_money(float(context.get("fixed_total") or 0))),
        ]
        if disc_pct:
            totals.append((f"Скидка {disc_pct:.0f}%", "—"))
        if tax_pct or context.get("tax_amount"):
            totals.append((f"Налог {tax_pct:.0f}%", _fmt_money(float(context.get("tax_amount") or 0))))
        totals.append(("ИТОГО", _fmt_money(grand)))
        pdf.pdf.ln(3)
        pdf.totals_table(totals, sum_col_w=sum_col_w)
        pdf.right(get_rubles_text(grand), bold=True)

    pdf.pdf.ln(4)
    pdf.line(
        "Полный юридический текст договора см. в версии Word. "
        "Этот PDF — спецификация и итоговая сумма для клиента."
    )
    _pdf_add_notes(pdf, context.get("tpl_footer_notes") or "")
    for lh in _pdf_letterhead_lines(context):
        pdf.line(lh)
    pdf.save(output_path)
    return output_path

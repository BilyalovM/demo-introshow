import docx

def update_template():
    doc = docx.Document('templates/contract_template.docx')
    
    # Update Paragraph 6 for based_on
    for p in doc.paragraphs:
        if 'Устава' in p.text and 'Арендатор' in p.text:
            # We want to replace "Устава" with "{{ based_on }}"
            for run in p.runs:
                if 'Устава' in run.text:
                    run.text = run.text.replace('Устава', '{{ based_on }}')
    
    # Update Table 1 (Арендатор)
    table_arendator = None
    for table in doc.tables:
        if len(table.rows) > 0 and len(table.columns) > 1:
            if 'Арендатор' in table.rows[0].cells[0].text or 'Арендатор' in table.rows[0].cells[1].text:
                table_arendator = table
                break
    
    if table_arendator:
        table_arendator.rows[2].cells[1].text = '{{ company_address }}'
        table_arendator.rows[4].cells[1].text = '{{ bank_name }}'
        table_arendator.rows[5].cells[1].text = '{{ kbe }}'
        table_arendator.rows[6].cells[1].text = '{{ bik }}'
        table_arendator.rows[7].cells[1].text = '{{ iban }}'
        
    # Update Table 3 (Equipment list)
    table_equipment = doc.tables[3]
    
    # Save the header row (row 0)
    # We will create a new table or just delete rows from the end
    
    # Delete all rows except the first one
    for i in range(len(table_equipment.rows) - 1, 0, -1):
        row = table_equipment.rows[i]
        tbl = table_equipment._tbl
        tbl.remove(row._tr)
        
    # Now we only have the header (row 0)
    # Add our loop row
    row_loop = table_equipment.add_row()
    row_loop.cells[0].text = '{% tr for item in items %}\n{{ item.name }}'
    row_loop.cells[1].text = '{{ item.quantity }}'
    row_loop.cells[2].text = '{{ item.price_text }}'
    row_loop.cells[3].text = '{{ item.days }}'
    row_loop.cells[4].text = '{{ item.subtotal_text }}\n{% tr endfor %}'
    
    # Add Итого row
    row_total = table_equipment.add_row()
    row_total.cells[0].text = ''
    row_total.cells[1].text = 'ИТОГО:'
    row_total.cells[2].text = ''
    row_total.cells[3].text = ''
    row_total.cells[4].text = '{{ equipment_total_text }}'
    
    # We also need an empty row or just style them
    # Let's save
    doc.save('templates/contract_template.docx')
    print("Template updated successfully.")

if __name__ == '__main__':
    update_template()

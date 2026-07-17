import os
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from num2words import num2words
from typing import Dict, Any

def get_rubles_text(amount: float) -> str:
    """
    Converts a number to Russian text.
    Handles the integer part and adds tiyn (decimals) for tenge.
    """
    try:
        integer_part = int(amount)
        decimal_part = int(round((amount - integer_part) * 100))
        
        text = num2words(integer_part, lang='ru')
        
        # format decimal part to always have 2 digits
        decimal_str = f"{decimal_part:02d}"
        
        return f"{text.capitalize()} тенге {decimal_str} тиын"
    except Exception as e:
        return f"{amount:,.2f} тенге"

def generate_contract(context: Dict[str, Any], template_path: str, output_path: str) -> str:
    """
    Generates a docx contract by injecting context into a template.
    Returns the path to the generated document.
    """
    doc = DocxTemplate(template_path)
    
    # Calculate the total in text
    total_cost = context.get('grand_total', 0.0)
    context['grand_total_text'] = get_rubles_text(total_cost)
    
    doc.render(context)
    doc.save(output_path)
    
    # Append appendix for photos and descriptions if they exist
    items = context.get('items', [])
    has_appendix = any(i.get('photo_url') or i.get('description') for i in items)
    
    if has_appendix:
        append_doc = Document(output_path)
        append_doc.add_page_break()
        append_doc.add_heading('Приложение: Спецификация оборудования', level=1)
        
        for item in items:
            if item.get('photo_url') or item.get('description'):
                append_doc.add_heading(item['name'], level=2)
                if item.get('description'):
                    append_doc.add_paragraph(item['description'])
                if item.get('photo_url'):
                    img_path = item['photo_url'].lstrip('/')
                    if os.path.exists(img_path):
                        append_doc.add_picture(img_path, width=Inches(3))
                        
        append_doc.save(output_path)
    
    return output_path

import docx
from docx.oxml import parse_xml

doc = docx.Document('templates/contract_template.docx')
table = doc.tables[3]

# Row 1 is the item row
tr = table.rows[1]._tr

for cell in table.rows[1].cells:
    for p in cell.paragraphs:
        p.text = p.text.replace('{% for item in items %}\n', '')
        p.text = p.text.replace('{% for item in items %}', '')
        p.text = p.text.replace('\n{% endfor %}', '')
        p.text = p.text.replace('{% endfor %}', '')

tr_for = parse_xml('<w:tr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:tc><w:p><w:r><w:t>{%tr for item in items %}</w:t></w:r></w:p></w:tc></w:tr>')
tr.addprevious(tr_for)

tr_endfor = parse_xml('<w:tr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:tc><w:p><w:r><w:t>{%tr endfor %}</w:t></w:r></w:p></w:tc></w:tr>')
tr.addnext(tr_endfor)

doc.save('templates/contract_template.docx')

import os
from document_generator import generate_contract
context = {
    "contract_number": "123",
    "contract_date": "10.10.2026",
    "company_name": "ООО Ромашка",
    "director_name": "Иванов И.И.",
    "iin_bin": "123456789",
    "iban": "KZ123",
    "based_on": "Устава",
    "company_address": "ул. Пушкина",
    "bank_name": "Kaspi",
    "kbe": "17",
    "bik": "CASP",
    "event_name": "Test Event",
    "event_date": "20.10.2026",
    "event_address": "Almaty",
    "items": [{"name": "Item 1", "quantity": 1, "price_text": "100", "days": 1, "subtotal_text": "100"}],
    "equipment_total_text": "100",
    "fixed_total": 0,
    "grand_total": 100,
    "grand_total_text": "Сто тенге",
    "discount_percentage": 0
}
generate_contract(context, "templates/contract_template.docx", "test_output_final.docx")
print("Render final successful.")
